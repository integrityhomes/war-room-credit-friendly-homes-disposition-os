from __future__ import annotations

from datetime import date
from io import BytesIO
from typing import Any, Iterable

import pandas as pd
from dateutil.relativedelta import relativedelta
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docxtpl import DocxTemplate

from .contract_insurance_control import apply_buyer_responsible_insurance_to_document

AMORTIZATION_MARKER = "[[AMORTIZATION_TABLE]]"


def format_currency(value: float) -> str:
    return f"${value:,.2f}"


def build_amortization_schedule(
    principal: float,
    annual_interest_rate: float,
    number_of_payments: int,
    first_payment_date: date,
    monthly_principal_interest: float,
    monthly_taxes: float,
    monthly_insurance: float,
) -> pd.DataFrame:
    rows: list[dict[str, object]] = []
    remaining_balance = principal
    monthly_interest_rate = annual_interest_rate / 100 / 12

    for payment_number in range(1, number_of_payments + 1):
        beginning_balance = remaining_balance
        interest_amount = beginning_balance * monthly_interest_rate
        principal_amount = monthly_principal_interest - interest_amount
        if payment_number == number_of_payments:
            principal_amount = beginning_balance
        principal_amount = min(max(principal_amount, 0.0), beginning_balance)
        actual_principal_interest = principal_amount + interest_amount
        ending_balance = max(beginning_balance - principal_amount, 0.0)
        payment_due_date = first_payment_date + relativedelta(months=payment_number - 1)
        total_payment = actual_principal_interest + monthly_taxes + monthly_insurance
        rows.append(
            {
                "Payment Number": payment_number,
                "Due Date": payment_due_date.strftime("%m/%d/%Y"),
                "Beginning Balance": round(beginning_balance, 2),
                "Principal and Interest": round(actual_principal_interest, 2),
                "Principal": round(principal_amount, 2),
                "Interest": round(interest_amount, 2),
                "Monthly Taxes": round(monthly_taxes, 2),
                "Monthly Insurance": round(monthly_insurance, 2),
                "Total Payment": round(total_payment, 2),
                "Ending Balance": round(ending_balance, 2),
            }
        )
        remaining_balance = ending_balance
        if remaining_balance <= 0:
            break
    return pd.DataFrame(rows)


def iter_all_paragraphs(document: Document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _set_paragraph_text_preserving_first_run(paragraph: Any, text: str) -> None:
    if paragraph.runs:
        first_run = paragraph.runs[0]
        first_run.text = text
        for extra_run in paragraph.runs[1:]:
            extra_run._element.getparent().remove(extra_run._element)
    else:
        paragraph.add_run(text)


def apply_residential_disclosure_answers(
    document: Document,
    yes_questions: Iterable[int],
    explanation_text: str,
) -> None:
    yes_question_numbers = {int(question_number) for question_number in yes_questions}
    inside_disclosure_section = False
    explanation_line_is_next = False

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        stripped_text = paragraph_text.strip()
        if stripped_text.startswith("RESIDENTIAL REAL PROPERTY DISCLOSURE REPORT"):
            inside_disclosure_section = True
            continue
        if not inside_disclosure_section:
            continue
        if stripped_text.startswith("Seller certifies that"):
            inside_disclosure_section = False
            continue
        if stripped_text.startswith("If any of the above are marked"):
            explanation_line_is_next = True
            continue
        if explanation_line_is_next and (stripped_text.startswith("_") or not stripped_text):
            replacement_text = explanation_text.strip() if explanation_text.strip() else "_" * 82
            _set_paragraph_text_preserving_first_run(paragraph, replacement_text)
            if paragraph.runs:
                paragraph.runs[0].font.name = "Times New Roman"
                paragraph.runs[0].font.size = Pt(10)
            explanation_line_is_next = False
            continue

        import re

        question_match = re.search(r"(?<!\d)([1-9]|1\d|2[0-4])\.\s", paragraph_text)
        if not question_match:
            continue
        question_number = int(question_match.group(1))
        question_statement = paragraph_text[question_match.start() :].strip()
        empty_box = "𝥷"
        check_mark = "✔️"
        if question_number in yes_question_numbers:
            answer_prefix = f"  {check_mark}         \t {empty_box}      \t   {empty_box}\t"
        else:
            answer_prefix = f"  {empty_box}         \t {check_mark}      \t   {empty_box}\t"
        _set_paragraph_text_preserving_first_run(paragraph, answer_prefix + question_statement)


def _set_repeat_table_header(row: Any) -> None:
    table_row_properties = row._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    table_row_properties.append(repeat_header)


def _set_cell_font_size(cell: Any, size: int = 8) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(size)


def _set_table_borders(table: Any) -> None:
    table_properties = table._tbl.tblPr
    existing_borders = table_properties.find(qn("w:tblBorders"))
    if existing_borders is not None:
        table_properties.remove(existing_borders)
    table_borders = OxmlElement("w:tblBorders")
    for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = OxmlElement(f"w:{edge_name}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), "000000")
        table_borders.append(edge)
    table_properties.append(table_borders)


def insert_amortization_table(document: Document, schedule: pd.DataFrame) -> None:
    marker_paragraph = next((paragraph for paragraph in iter_all_paragraphs(document) if AMORTIZATION_MARKER in paragraph.text), None)
    if marker_paragraph is None:
        raise ValueError("The amortization marker was not found in the Word template.")
    marker_paragraph.text = ""
    table = document.add_table(rows=1, cols=6)
    _set_table_borders(table)
    headers = ["Pmt #", "Due Date", "Payment", "Principal", "Interest", "Balance"]
    header_row = table.rows[0]
    _set_repeat_table_header(header_row)
    for index, header in enumerate(headers):
        header_row.cells[index].text = header
        _set_cell_font_size(header_row.cells[index], 8)
    for _, payment in schedule.iterrows():
        row_cells = table.add_row().cells
        values = [
            str(int(payment["Payment Number"])),
            str(payment["Due Date"]),
            format_currency(float(payment["Principal and Interest"])),
            format_currency(float(payment["Principal"])),
            format_currency(float(payment["Interest"])),
            format_currency(float(payment["Ending Balance"])),
        ]
        for index, value in enumerate(values):
            row_cells[index].text = value
            _set_cell_font_size(row_cells[index], 8)
    marker_paragraph._p.addnext(table._tbl)


def generate_contract_document(
    *,
    template_bytes: bytes,
    context: dict[str, Any],
    amortization_schedule: pd.DataFrame,
    disclosure_yes_questions: Iterable[int] = (),
    disclosure_explanation: str = "",
) -> tuple[bytes, tuple[str, ...]]:
    """Render an approved template using the reviewed V14 operation order."""
    if not template_bytes:
        raise ValueError("Approved contract template bytes are required.")

    template = DocxTemplate(BytesIO(template_bytes))
    template.render(context, autoescape=True)
    rendered_buffer = BytesIO()
    template.save(rendered_buffer)
    rendered_buffer.seek(0)

    document = Document(rendered_buffer)
    insurance_changes = tuple(apply_buyer_responsible_insurance_to_document(document=document, context=context))
    apply_residential_disclosure_answers(
        document=document,
        yes_questions=disclosure_yes_questions,
        explanation_text=disclosure_explanation,
    )
    insert_amortization_table(document, amortization_schedule)

    completed_buffer = BytesIO()
    document.save(completed_buffer)
    completed_buffer.seek(0)
    return completed_buffer.getvalue(), insurance_changes
