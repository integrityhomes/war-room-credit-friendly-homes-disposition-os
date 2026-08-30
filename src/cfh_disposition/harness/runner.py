from __future__ import annotations

import argparse
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document

from cfh_disposition.commandcore_offer_engine import OfferDealInput, analyze_deal
from cfh_disposition.contract_generation_pipeline import generate_and_store_contract

from .fixtures import FIXTURE_FAMILY, FIXTURE_SOURCE, load_fixture_family
from .mode import HarnessMode, parse_mode
from .report import HarnessReport, write_report
from .side_effects import ActionType, SideEffectBus

DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class _FakeBucket:
    def __init__(self, template_bytes: bytes) -> None:
        self.template_bytes = template_bytes
        self.uploads: list[dict[str, Any]] = []

    def download(self, path: str) -> bytes:
        if not path.startswith("fixtures/templates/"):
            raise AssertionError("Harness storage may only read fixture templates.")
        return self.template_bytes

    def upload(self, **kwargs: Any) -> dict[str, str]:
        self.uploads.append(kwargs)
        return {"path": str(kwargs["path"])}


class _FakeStorage:
    def __init__(self, template_bytes: bytes) -> None:
        self.bucket = _FakeBucket(template_bytes)

    def get_bucket(self, name: str) -> dict[str, bool]:
        if name != "commandcore-contract-documents":
            raise AssertionError("Harness contract storage must use the existing private bucket name.")
        return {"public": False}

    def from_(self, name: str) -> _FakeBucket:
        if name != "commandcore-contract-documents":
            raise AssertionError("Harness contract storage must use the existing private bucket name.")
        return self.bucket


class _FakeClient:
    def __init__(self, template_bytes: bytes) -> None:
        self.storage = _FakeStorage(template_bytes)


def _fixture_template_bytes() -> bytes:
    document = Document()
    document.add_paragraph("HARNESS FIXTURE - Illinois Contract for Deed - {{ BUYER_NAMES }}")
    document.add_paragraph("Property: {{ PROPERTY_ADDRESS }}")
    document.add_paragraph("Seller: {{ SELLER_NAME }}")
    document.add_paragraph("Purchase price: {{ SALES_PRICE }}")
    document.add_paragraph("Insurance: Seller shall maintain insurance under the approved included-insurance template.")
    document.add_paragraph("[[AMORTIZATION_TABLE]]")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _offer_input(fixture: dict[str, Any]) -> OfferDealInput:
    deal = fixture["deal"]
    property_record = fixture["property"]
    return OfferDealInput(
        address=str(property_record["address"]),
        market=str(deal["market"]),
        lead_type=str(deal["lead_type"]),
        exit_mode=str(deal["exit_mode"]),
        asking_price=float(deal["asking_price"]),
        rent=float(deal["rent"]),
        beds=float(deal["beds"]),
        baths=float(deal["baths"]),
        sqft=float(deal["sqft"]),
        taxes=float(deal["taxes"]),
        status=str(deal["status"]),
        occupancy=str(deal["occupancy"]),
        livable=str(deal["livable"]),
        days_on_market=int(deal["days_on_market"]),
        notes=str(deal["notes"]),
        arv=float(deal["arv"]),
        repairs=float(deal["repairs"]),
        rent_source=str(deal["rent_source"]),
        rent_confidence=str(deal["rent_confidence"]),
        rent_verification_needed=str(deal["rent_verification_needed"]),
    )


def _facts_document(fixture: dict[str, Any]) -> dict[str, Any]:
    deal = fixture["deal"]
    contact = fixture["contact"]
    property_record = fixture["property"]
    return {
        "id": "FIXTURE-CONTRACT-FACTS-HARRIS-0001",
        "document_type": "contract_prep_facts",
        "contract_type": deal["contract_type"],
        "internal_only": True,
        "fixture_source": FIXTURE_SOURCE,
        "facts": {
            "contract_type": deal["contract_type"],
            "state": property_record["state"],
            "seller_name": contact["name"],
            "seller_mailing_address": "100 TEST Seller Way, Chesapeake, VA 23320",
            "seller_formation_state": "VA",
            "property_address": property_record["address"],
            "property_county": property_record["county"],
            "legal_description": property_record["legal_description"],
            "parcel_number": property_record["parcel_number"],
            "buyer_1_name": deal["buyer_1_name"],
            "purchase_price": deal["purchase_price"],
            "down_payment": deal["down_payment"],
            "interest_rate": deal["interest_rate"],
            "number_of_payments": deal["number_of_payments"],
            "first_payment_date": deal["first_payment_date"],
            "monthly_taxes": deal["monthly_taxes"],
            "monthly_insurance": deal["monthly_insurance"],
            "insurance_included": deal["insurance_included"],
            "contract_date": deal["contract_date"],
            "payment_payee": deal["payment_payee"],
            "payment_address": deal["payment_address"],
            "payment_system": deal["payment_system"],
            "current_lien_disclosure": "Not Applicable",
            "disclosure_yes_questions": [],
            "disclosure_explanation": "",
        },
        "links": {"deal_id": deal["id"]},
    }


def _approved_fixture_template() -> dict[str, Any]:
    return {
        "id": "FIXTURE-TEMPLATE-IL-CFD-0001",
        "name": "HARNESS-Illinois-CFD-v1.docx",
        "document_type": "contract_template",
        "contract_type": "Illinois Contract for Deed",
        "state": "IL",
        "version": 1,
        "template_version": "fixture-v1",
        "status": "approved",
        "legal_review_status": "approved",
        "legal_approved": True,
        "approved_for_use": True,
        "storage_object_path": "fixtures/templates/il/contract-for-deed/v1/template.docx",
        "internal_only": True,
        "fixture_source": FIXTURE_SOURCE,
    }


def run_offer_no_send(mode: str | HarnessMode | None = None) -> HarnessReport:
    fixture = load_fixture_family()
    deal = fixture["deal"]
    bus = SideEffectBus(mode)
    analysis = analyze_deal(_offer_input(fixture))
    offer = {
        **fixture["offer"],
        "starting_offer": analysis["best"]["first_offer"],
        "max_offer": analysis["best"]["max_offer"],
        "recommended_exit": analysis["best_exit"],
        "grade": analysis["grade"],
        "internal_only": True,
        "external_action_started": False,
    }
    bus.request(
        ActionType.CRM_COMMIT,
        {"record_type": "offer", "record_id": offer["id"], "record": offer},
        deal=deal,
        owner_approval=fixture["approval"],
    )
    send = bus.request(
        ActionType.OFFER_SEND,
        {
            "to": fixture["contact"]["email"],
            "deal_id": deal["id"],
            "offer_amount": offer["starting_offer"],
            "max_offer": offer["max_offer"],
        },
        deal=deal,
        owner_approval=fixture["approval"],
    )
    passed = send.decision == "blocked" and bus.provider_calls == 0
    return HarnessReport(
        scenario="offer_no_send",
        mode=bus.mode.value,
        fixture_family=FIXTURE_FAMILY,
        verdict="PASS" if passed else "FAIL",
        provider_calls=bus.provider_calls,
        actions=list(bus.records),
        artifacts={"offer": offer, "offer_analysis": analysis},
    )


def run_contract_no_sign(mode: str | HarnessMode | None = None) -> HarnessReport:
    fixture = load_fixture_family()
    deal = fixture["deal"]
    bus = SideEffectBus(mode)
    fake_client = _FakeClient(_fixture_template_bytes())
    generated = generate_and_store_contract(
        client=fake_client,
        deal_id=deal["id"],
        facts_document=_facts_document(fixture),
        all_documents=[_approved_fixture_template(), fixture["contract_draft"]],
    )
    document_record = {
        **generated.document_record,
        "id": "FIXTURE-CONTRACT-HARRIS-V2",
        "internal_only": True,
        "external_action_started": False,
        "fixture_source": FIXTURE_SOURCE,
    }
    contract_send = bus.request(
        ActionType.CONTRACT_SEND,
        {"document_id": document_record["id"], "document_version": document_record["version"], "to": fixture["contact"]["email"]},
        deal=deal,
        owner_approval=fixture["approval"],
    )
    contract_sign = bus.request(
        ActionType.CONTRACT_SIGN,
        {"document_id": document_record["id"], "document_version": document_record["version"]},
        deal=deal,
        owner_approval=fixture["approval"],
    )
    passed = (
        document_record["version"] == 2
        and contract_send.decision == "blocked"
        and contract_sign.decision == "blocked"
        and bus.provider_calls == 0
    )
    return HarnessReport(
        scenario="contract_no_sign",
        mode=bus.mode.value,
        fixture_family=FIXTURE_FAMILY,
        verdict="PASS" if passed else "FAIL",
        provider_calls=bus.provider_calls,
        actions=list(bus.records),
        artifacts={
            "generated_document": {
                "id": document_record["id"],
                "version": document_record["version"],
                "storage_bucket": document_record["storage_bucket"],
                "storage_object_path": document_record["storage_object_path"],
                "internal_only": document_record["internal_only"],
                "external_action_started": document_record["external_action_started"],
            },
            "fake_private_upload_count": len(fake_client.storage.bucket.uploads),
        },
    )


def run_scenario(scenario: str, mode: str | HarnessMode | None = None) -> HarnessReport:
    if scenario == "offer_no_send":
        return run_offer_no_send(mode)
    if scenario == "contract_no_sign":
        return run_contract_no_sign(mode)
    raise ValueError(f"Unknown harness scenario: {scenario!r}")


def main() -> int:
    parser = argparse.ArgumentParser(description="Run a CommandCore Test & Simulation Harness scenario.")
    parser.add_argument("--scenario", choices=("offer_no_send", "contract_no_sign"), required=True)
    parser.add_argument("--mode", choices=tuple(mode.value for mode in HarnessMode), default=HarnessMode.SIMULATION.value)
    parser.add_argument("--output-dir", default="artifacts")
    args = parser.parse_args()

    report = run_scenario(args.scenario, parse_mode(args.mode))
    json_path, markdown_path = write_report(report, Path(args.output_dir))
    print(report.markdown(), end="")
    print(f"JSON report: {json_path}")
    print(f"Markdown report: {markdown_path}")
    return 0 if report.verdict == "PASS" else 1


if __name__ == "__main__":
    raise SystemExit(main())
