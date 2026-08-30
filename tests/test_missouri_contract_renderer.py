from __future__ import annotations

from datetime import date
from io import BytesIO

from docx import Document

from cfh_disposition.contract_document_renderer import build_amortization_schedule
from cfh_disposition.missouri_contract_renderer import (
    build_missouri_contract_context,
    generate_missouri_contract_document,
    ordinal_day,
)


def _context() -> dict:
    return build_missouri_contract_context(
        property_address="123 Main St, Saint Louis, MO 63101",
        seller_name="Example Seller LLC",
        seller_address="456 Seller Rd, Saint Louis, MO 63102",
        buyer_1_name="Buyer One",
        buyer_2_name="Buyer Two",
        buyer_1_email="one@example.com",
        buyer_2_email="two@example.com",
        buyer_1_phone="555-1111",
        buyer_2_phone="555-2222",
        contract_date=date(2026, 8, 30),
        first_payment_date=date(2026, 10, 1),
        sales_price=50000,
        down_payment=2500,
        amount_financed=47500,
        annual_interest_rate=10,
        number_of_payments=360,
        monthly_principal_interest=416.79,
        monthly_taxes=75,
        monthly_servicing_fee=25,
        total_monthly_payment=516.79,
        conversion_rent=900,
        payment_payee="Example Seller LLC",
        payment_address="456 Seller Rd, Saint Louis, MO 63102",
        payment_system="Buildium property management website",
        late_fee_percent=10,
        grace_period_days=5,
        apr=10,
        finance_charge=102544.4,
        total_of_payments=150044.4,
        use_primary=False,
        use_investment=True,
        use_fix_flip=False,
        use_family=False,
        use_short_term=False,
        use_landlord=True,
        use_other=False,
        use_other_text="",
    )


def test_missouri_context_matches_v14_fields() -> None:
    context = _context()
    assert context["PROPERTY_CITY"] == "Saint Louis"
    assert context["CONTRACT_DAY"] == "30th"
    assert context["CONTRACT_MONTH_YEAR"] == "August 2026"
    assert context["AMORTIZATION_END_DATE"] == "September 1, 2056"
    assert context["LOAN_TERM"] == "360 months (30 years)"
    assert context["MONTHLY_SERVICING_FEE"] == "$25.00"
    assert context["CONVERSION_RENT"] == "$900.00"
    assert context["USE_INVESTMENT"] == "[X]"
    assert context["USE_LANDLORD"] == "[X]"
    assert context["USE_PRIMARY"] == "[ ]"


def test_ordinal_day_handles_teens_and_common_suffixes() -> None:
    assert ordinal_day(1) == "1st"
    assert ordinal_day(2) == "2nd"
    assert ordinal_day(3) == "3rd"
    assert ordinal_day(11) == "11th"
    assert ordinal_day(22) == "22nd"


def test_missouri_renderer_renders_template_and_inserts_amortization() -> None:
    source = Document()
    source.add_paragraph("Agreement for {{ BUYER_NAMES }} at {{ PROPERTY_ADDRESS }}")
    source.add_paragraph("{{ AMORTIZATION_SCHEDULE }}")
    template_buffer = BytesIO()
    source.save(template_buffer)

    schedule = build_amortization_schedule(
        principal=1000,
        annual_interest_rate=0,
        number_of_payments=2,
        first_payment_date=date(2026, 10, 1),
        monthly_principal_interest=500,
        monthly_taxes=0,
        monthly_insurance=0,
    )
    rendered = generate_missouri_contract_document(
        template_bytes=template_buffer.getvalue(),
        context=_context(),
        amortization_schedule=schedule,
    )

    completed = Document(BytesIO(rendered))
    text = "\n".join(paragraph.text for paragraph in completed.paragraphs)
    assert "Buyer One and Buyer Two" in text
    assert "123 Main St, Saint Louis, MO 63101" in text
    assert len(completed.tables) == 1
    assert completed.tables[0].rows[0].cells[0].text == "Pmt #"
    assert completed.tables[0].rows[1].cells[0].text == "1"
