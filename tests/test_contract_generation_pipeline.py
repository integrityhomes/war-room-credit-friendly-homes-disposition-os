from __future__ import annotations

from datetime import date
from io import BytesIO

import pytest
from docx import Document

from cfh_disposition.contract_generation_pipeline import (
    ContractGenerationError,
    earliest_illinois_execution_date,
    generate_and_store_contract,
    select_exact_approved_template,
)


DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _template_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Illinois Contract for Deed - {{ BUYER_NAMES }}")
    document.add_paragraph("Property: {{ PROPERTY_ADDRESS }}")
    document.add_paragraph("Seller: {{ SELLER_NAME }}")
    document.add_paragraph("Purchase price: {{ SALES_PRICE }}")
    document.add_paragraph("Insurance: Seller shall maintain insurance under the approved included-insurance template.")
    document.add_paragraph("[[AMORTIZATION_TABLE]]")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class FakeBucket:
    def __init__(self, template_bytes: bytes) -> None:
        self.template_bytes = template_bytes
        self.uploads: list[dict] = []

    def download(self, path: str) -> bytes:
        assert path.startswith("templates/")
        return self.template_bytes

    def upload(self, **kwargs):  # noqa: ANN003, ANN201
        self.uploads.append(kwargs)
        return {"path": kwargs["path"]}


class FakeStorage:
    def __init__(self, template_bytes: bytes) -> None:
        self.bucket = FakeBucket(template_bytes)

    def get_bucket(self, name: str) -> dict[str, bool]:
        assert name == "commandcore-contract-documents"
        return {"public": False}

    def from_(self, name: str) -> FakeBucket:
        assert name == "commandcore-contract-documents"
        return self.bucket


class FakeClient:
    def __init__(self, template_bytes: bytes) -> None:
        self.storage = FakeStorage(template_bytes)


def _approved_template(*, template_id: str = "template-2", version: int = 2, state: str = "IL") -> dict:
    return {
        "id": template_id,
        "name": f"Illinois-CFD-v{version}.docx",
        "document_type": "contract_template",
        "contract_type": "Illinois Contract for Deed",
        "state": state,
        "version": version,
        "template_version": f"v{version}",
        "status": "approved",
        "legal_review_status": "approved",
        "legal_approved": True,
        "approved_for_use": True,
        "storage_object_path": f"templates/{state.lower()}/illinois-contract-for-deed/v{version}/template.docx",
    }


def _facts_document() -> dict:
    return {
        "id": "facts-9",
        "document_type": "contract_prep_facts",
        "contract_type": "Illinois Contract for Deed",
        "facts": {
            "contract_type": "Illinois Contract for Deed",
            "state": "IL",
            "seller_name": "Seller LLC",
            "seller_mailing_address": "100 Main St, Chesapeake, VA 23320",
            "seller_formation_state": "VA",
            "property_address": "330 N 24th St, Decatur, IL 62521",
            "property_county": "Macon",
            "legal_description": "Lot 1 in Test Subdivision",
            "parcel_number": "04-12-34-567-890",
            "assessed_value": "$18,000.00",
            "fair_cash_value": "$54,000.00",
            "last_tax_bill": "$1,200.00",
            "buyer_1_name": "Test Buyer",
            "buyer_2_name": "",
            "purchase_price": "100000",
            "down_payment": "5000",
            "interest_rate": "10",
            "number_of_payments": "360",
            "first_payment_date": "2026-10-01",
            "monthly_taxes": "100",
            "monthly_insurance": "75",
            "insurance_included": "Yes",
            "contract_date": "2026-08-31",
            "payment_payee": "Seller LLC",
            "payment_address": "100 Main St, Chesapeake, VA 23320",
            "payment_system": "Approved servicing portal",
            "current_lien_disclosure": "Not Applicable",
            "disclosure_yes_questions": [],
            "disclosure_explanation": "",
        },
        "links": {"deal_id": "deal-123"},
    }


def test_exact_approved_template_prefers_latest_matching_version() -> None:
    documents = [
        _approved_template(template_id="older", version=1),
        _approved_template(template_id="newer", version=3),
        _approved_template(template_id="wrong-state", version=99, state="MO"),
        {**_approved_template(template_id="pending", version=100), "status": "needs_legal_approval", "approved_for_use": False},
    ]
    selected = select_exact_approved_template(
        documents,
        contract_type="Illinois Contract for Deed",
        state="IL",
    )
    assert selected["id"] == "newer"


def test_template_selection_fails_closed_without_exact_approval() -> None:
    with pytest.raises(ContractGenerationError, match="exactly matches"):
        select_exact_approved_template(
            [_approved_template(state="MO")],
            contract_type="Illinois Contract for Deed",
            state="IL",
        )


def test_execution_date_skips_weekend_and_labor_day() -> None:
    # Notice Friday before Labor Day 2026: Tue/Wed/Thu are the three full business days; execution is Friday.
    assert earliest_illinois_execution_date(date(2026, 9, 4)) == date(2026, 9, 11)


def test_generate_and_store_contract_is_private_versioned_and_auditable() -> None:
    client = FakeClient(_template_bytes())
    documents = [
        _approved_template(version=1, template_id="template-1"),
        _approved_template(version=2, template_id="template-2"),
        {
            "id": "generated-old",
            "document_type": "generated_contract",
            "contract_type": "Illinois Contract for Deed",
            "version": 1,
            "links": {"deal_id": "deal-123"},
        },
    ]

    generated = generate_and_store_contract(
        client=client,
        deal_id="deal-123",
        facts_document=_facts_document(),
        all_documents=documents,
    )

    record = generated.document_record
    assert record["document_type"] == "generated_contract"
    assert record["version"] == 2
    assert record["status"] == "generated_needs_review"
    assert record["approved_legal_template_id"] == "template-2"
    assert record["source_facts_document_id"] == "facts-9"
    assert record["insurance_version"] == "Insurance Included"
    assert record["document_assembled"] is True
    assert record["legal_terms_generated"] is False
    assert record["legal_terms_changed_by_commandcore"] is False
    assert record["signing_started"] is False
    assert record["external_action_started"] is False
    assert record["generation_provenance"]["template_version"] == 2
    assert record["generation_provenance"]["immutable_version"] == 2
    assert record["storage_bucket"] == "commandcore-contract-documents"
    assert "/generated_contract/v2/" in record["storage_object_path"]
    assert client.storage.bucket.uploads[0]["file_options"]["upsert"] == "false"

    output = Document(BytesIO(generated.generated_bytes))
    output_text = "\n".join(paragraph.text for paragraph in output.paragraphs)
    assert "Test Buyer" in output_text
    assert "330 N 24th St" in output_text
    assert "$100,000.00" in output_text
    assert output.tables
    assert generated.activity_record["activity_type"] == "contract_generated"
    assert generated.activity_record["details"]["external_action_started"] is False
