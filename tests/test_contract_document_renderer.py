from __future__ import annotations

from datetime import date
from io import BytesIO

from docx import Document

from cfh_disposition.contract_document_renderer import build_amortization_schedule, generate_contract_document
from cfh_disposition.contract_insurance_control import (
    INSURANCE_VERSION_BUYER_RESPONSIBLE,
    INSURANCE_VERSION_INCLUDED,
    buyer_responsible_version_ready,
    insurance_version_label,
    normalize_insurance_included,
)


def _template_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Contract for {{ BUYER_NAMES }}")
    document.add_paragraph("4. Buyer shall pay to Seller the total sum of {{ SALES_PRICE }}")
    document.add_paragraph("The last hazard insurance bill for the premises was {{ LAST_INSURANCE_BILL }}")
    document.add_paragraph("12. A. Buyer shall pay insurance when due")
    document.add_paragraph("B. Seller shall be obligated to pay taxes and insurance when due")
    document.add_paragraph("payment to Escrow Agent who shall add the same")
    document.add_paragraph("(Upon written notice by Seller to Buyer and Escrow Agent")
    document.add_paragraph("increased by 1/12 of any increase in taxes and insurance")
    document.add_paragraph("adjustment.) (Buyer agrees to pay said sums to Escrow Agent")
    document.add_paragraph("notification of the same by Seller.)")
    document.add_paragraph("14. For the Duration of this Agreement, Seller shall keep all improvements insured")
    document.add_paragraph("shall insure all of Buyer’s Items of personal property")
    document.add_paragraph("damaged by casualty, Buyer shall promptly cause all improvements")
    document.add_paragraph("Buyer shall fail to cause such repairs or rebuilding")
    document.add_paragraph("insurance proceeds shall be instead paid to Escrow Agent")
    document.add_paragraph("balance of any mortgage on Premises")
    document.add_paragraph("balance, if any, shall then be paid to the Buyer.")
    document.add_paragraph("The Buyer agrees to pay any insurance deductible")
    document.add_paragraph("If the proposed transaction is canceled or the Contract is not fully executed on {{ EARLIEST_EXECUTION_DATE }}")
    document.add_paragraph("RESIDENTIAL REAL PROPERTY DISCLOSURE REPORT")
    document.add_paragraph("1. Seller occupied the property within the last 12 months")
    document.add_paragraph("If any of the above are marked Yes, explain:")
    document.add_paragraph("________________________________________________________________")
    document.add_paragraph("Seller certifies that the information is true")
    document.add_paragraph("[[AMORTIZATION_TABLE]]")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _context(status: str) -> dict[str, object]:
    return {
        "BUYER_NAMES": "Test Buyer",
        "SALES_PRICE": "$100,000.00",
        "DOWN_PAYMENT": "$5,000.00",
        "INTEREST_START_DATE": "January 1, 2027",
        "INTEREST_RATE": "10.0000%",
        "TOTAL_MONTHLY_PAYMENT": "$1,050.00",
        "MONTHLY_PRINCIPAL_INTEREST": "$950.00",
        "MONTHLY_TAXES": "$100.00",
        "FIRST_PAYMENT_DATE": "February 1, 2027",
        "PAYMENT_PAYEE": "Seller LLC",
        "PAYMENT_ADDRESS": "100 Main St",
        "PAYMENT_SYSTEM": "approved servicing system",
        "LAST_INSURANCE_BILL": "$900.00",
        "EARLIEST_EXECUTION_DATE": "January 8, 2027",
        "_INSURANCE_STATUS": status,
        "_PRIOR_YEAR_INSURANCE_KNOWN": False,
    }


def _document_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            paragraphs.extend(cell.text for cell in row.cells)
    return "\n".join(paragraphs)


def test_insurance_status_is_fail_closed_and_versioned() -> None:
    assert normalize_insurance_included("YES") == "yes"
    assert normalize_insurance_included("No") == "no"
    assert insurance_version_label("yes") == INSURANCE_VERSION_INCLUDED
    assert insurance_version_label("no") == INSURANCE_VERSION_BUYER_RESPONSIBLE
    assert buyer_responsible_version_ready() is True


def test_yes_version_keeps_existing_seller_insurance_template_language() -> None:
    schedule = build_amortization_schedule(95_000, 10, 2, date(2027, 2, 1), 48_000, 100, 75)
    output, changes = generate_contract_document(
        template_bytes=_template_bytes(),
        context=_context("yes"),
        amortization_schedule=schedule,
        disclosure_yes_questions=(1,),
        disclosure_explanation="Seller occupied the property.",
    )
    text = _document_text(output)
    assert changes == ()
    assert "Seller shall be obligated to pay taxes and insurance when due" in text
    assert "Seller shall keep all improvements insured" in text
    assert "Pmt #" in text
    assert "Seller occupied the property." in text


def test_no_version_applies_every_locked_buyer_responsible_clause() -> None:
    schedule = build_amortization_schedule(95_000, 10, 2, date(2027, 2, 1), 48_000, 100, 0)
    output, changes = generate_contract_document(
        template_bytes=_template_bytes(),
        context=_context("no"),
        amortization_schedule=schedule,
    )
    text = _document_text(output)
    assert set(changes) == {
        "monthly_payment",
        "prior_year_insurance",
        "insurance_responsibility_and_tax",
        "hazard_and_casualty",
        "temporary_early_possession",
    }
    assert "Insurance is not included in the Installment Payment" in text
    assert "additional insured" in text
    assert "Five Hundred Thousand Dollars ($500,000)" in text
    assert "Seller shall be obligated to pay taxes and insurance when due" not in text
    assert "Seller shall keep all improvements insured" not in text
    assert "prior-year insurance premium was zero" in text
